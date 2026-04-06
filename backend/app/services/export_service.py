from io import BytesIO

from docx import Document

from app.schemas.resume import ResumeRead


def build_docx_bytes(resume: ResumeRead) -> bytes:
  doc = Document()
  doc.add_heading(resume.content.full_name or resume.title, level=1)

  if resume.content.email or resume.content.phone:
    doc.add_paragraph(f"{resume.content.email} | {resume.content.phone}".strip(" |"))

  doc.add_heading("Summary", level=2)
  doc.add_paragraph(resume.content.summary or "N/A")

  doc.add_heading("Skills", level=2)
  if resume.content.skills:
    doc.add_paragraph(", ".join(resume.content.skills))
  else:
    doc.add_paragraph("N/A")

  doc.add_heading("Experience", level=2)
  if resume.content.experience:
    for item in resume.content.experience:
      doc.add_paragraph(item, style="List Bullet")
  else:
    doc.add_paragraph("N/A")

  doc.add_heading("Education", level=2)
  if resume.content.education:
    for item in resume.content.education:
      doc.add_paragraph(item, style="List Bullet")
  else:
    doc.add_paragraph("N/A")

  buffer = BytesIO()
  doc.save(buffer)
  buffer.seek(0)
  return buffer.read()


def build_pdf_bytes(resume: ResumeRead) -> bytes:
  try:
    from weasyprint import HTML
  except Exception as exc:
    raise RuntimeError(
      "WeasyPrint runtime dependencies are missing. Install GTK/Pango libraries on Windows."
    ) from exc

  skills = "".join(f"<li>{skill}</li>" for skill in resume.content.skills) or "<li>N/A</li>"
  experience = "".join(f"<li>{item}</li>" for item in resume.content.experience) or "<li>N/A</li>"
  education = "".join(f"<li>{item}</li>" for item in resume.content.education) or "<li>N/A</li>"

  html = f"""
  <html>
    <head>
      <style>
        body {{ font-family: Arial, sans-serif; margin: 32px; color: #222; }}
        h1 {{ margin-bottom: 0; }}
        .meta {{ color: #555; margin-top: 8px; }}
        h2 {{ margin-top: 26px; border-bottom: 1px solid #ddd; padding-bottom: 6px; }}
        ul {{ margin-top: 10px; }}
      </style>
    </head>
    <body>
      <h1>{resume.content.full_name or resume.title}</h1>
      <div class="meta">{resume.content.email} | {resume.content.phone}</div>
      <h2>Summary</h2>
      <p>{resume.content.summary or "N/A"}</p>
      <h2>Skills</h2>
      <ul>{skills}</ul>
      <h2>Experience</h2>
      <ul>{experience}</ul>
      <h2>Education</h2>
      <ul>{education}</ul>
    </body>
  </html>
  """
  return HTML(string=html).write_pdf()
