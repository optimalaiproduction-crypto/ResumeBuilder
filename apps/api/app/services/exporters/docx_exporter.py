from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor

NAME_SIZE = 22
SECTION_HEADING_SIZE = 13
SUBHEADING_SIZE = 11.5
BODY_SIZE = 10.5
META_SIZE = 10


def _text(value: object) -> str:
  return str(value or "").strip()


def _content_spacing(
  value: str,
  *,
  short: float,
  medium: float,
  long: float,
  short_limit: int = 70,
  medium_limit: int = 170,
) -> float:
  length = len(_text(value))
  if length <= short_limit:
    return short
  if length <= medium_limit:
    return medium
  return long


def _clean_unique(items: list[object]) -> list[str]:
  seen: set[str] = set()
  clean: list[str] = []
  for item in items:
    value = _text(item)
    if not value:
      continue
    key = value.casefold()
    if key in seen:
      continue
    seen.add(key)
    clean.append(value)
  return clean


def _set_document_defaults(doc: Document) -> None:
  section = doc.sections[0]
  section.top_margin = Inches(0.65)
  section.bottom_margin = Inches(0.65)
  section.left_margin = Inches(0.7)
  section.right_margin = Inches(0.7)

  normal_style = doc.styles["Normal"]
  normal_style.font.name = "Calibri"
  normal_style.font.size = Pt(BODY_SIZE)

  try:
    bullet_style = doc.styles["List Bullet"]
  except KeyError:
    bullet_style = None
  if bullet_style is not None:
    bullet_style.font.name = "Calibri"
    bullet_style.font.size = Pt(BODY_SIZE)


def _add_paragraph(
  doc: Document,
  text: str,
  *,
  size: float = 11,
  bold: bool = False,
  italic: bool = False,
  color: tuple[int, int, int] = (17, 24, 39),
  space_before: float = 0,
  space_after: float = 0,
) -> None:
  paragraph = doc.add_paragraph()
  paragraph.paragraph_format.space_before = Pt(space_before)
  paragraph.paragraph_format.space_after = Pt(space_after)
  run = paragraph.add_run(text)
  run.bold = bold
  run.italic = italic
  run.font.size = Pt(size)
  run.font.color.rgb = RGBColor(*color)


def _add_section_heading(doc: Document, label: str) -> None:
  _add_paragraph(
    doc,
    label.upper(),
    size=SECTION_HEADING_SIZE,
    bold=True,
    color=(31, 41, 55),
    space_before=8,
    space_after=2.5,
  )


def _add_subheading(doc: Document, label: str, *, space_after: float = 1.0) -> None:
  _add_paragraph(
    doc,
    label,
    size=SUBHEADING_SIZE,
    bold=True,
    color=(17, 24, 39),
    space_after=space_after,
  )


def export_docx(title: str, content: dict) -> bytes:
  basics = content.get("basics", {})
  full_name = _text(basics.get("fullName")) or _text(title) or "Resume"
  headline = _text(basics.get("title"))
  summary = _text(content.get("summary"))
  skills = _clean_unique(content.get("skills", []))

  work_items = []
  for item in content.get("workExperience", []):
    role = _text(item.get("title"))
    company = _text(item.get("company"))
    start = _text(item.get("startDate"))
    end = "Present" if item.get("current") else _text(item.get("endDate"))
    bullets = _clean_unique(item.get("bullets", []))
    if role or company or start or end or bullets:
      work_items.append(
        {
          "role": role,
          "company": company,
          "start": start,
          "end": end,
          "bullets": bullets
        }
      )

  education_items = []
  for item in content.get("education", []):
    degree = _text(item.get("degree"))
    institution = _text(item.get("institution"))
    if degree or institution:
      education_items.append({"degree": degree, "institution": institution})

  project_items = []
  for item in content.get("projects", []):
    name = _text(item.get("name"))
    description = _text(item.get("description"))
    link = _text(item.get("link"))
    bullets = _clean_unique(item.get("bullets", []))
    if name or description or link or bullets:
      project_items.append(
        {
          "name": name,
          "description": description,
          "link": link,
          "bullets": bullets
        }
      )

  certification_items = []
  for item in content.get("certifications", []):
    name = _text(item.get("name"))
    issuer = _text(item.get("issuer"))
    cert_date = _text(item.get("date"))
    link = _text(item.get("link"))
    if name or issuer or cert_date or link:
      certification_items.append(
        {
          "name": name,
          "issuer": issuer,
          "date": cert_date,
          "link": link
        }
      )

  doc = Document()
  _set_document_defaults(doc)
  _add_paragraph(doc, full_name, size=NAME_SIZE, bold=True, space_after=1)
  if headline and headline.casefold() != full_name.casefold():
    _add_subheading(doc, headline, space_after=1.6)

  contact = " | ".join(
    [
      _text(basics.get("email")),
      _text(basics.get("phone")),
      _text(basics.get("location")),
      _text(basics.get("linkedin")),
      _text(basics.get("website"))
    ]
  ).strip(" |")
  if contact:
    _add_paragraph(
      doc,
      contact,
      size=META_SIZE,
      color=(75, 85, 99),
      space_after=_content_spacing(contact, short=4.0, medium=5.0, long=6.0),
    )

  if summary:
    _add_section_heading(doc, "Summary")
    _add_paragraph(
      doc,
      summary,
      size=BODY_SIZE,
      color=(17, 24, 39),
      space_after=_content_spacing(summary, short=1.6, medium=2.2, long=2.8),
    )

  if skills:
    _add_section_heading(doc, "Skills")
    skills_text = ", ".join(skills)
    _add_paragraph(
      doc,
      skills_text,
      size=BODY_SIZE,
      color=(17, 24, 39),
      space_after=_content_spacing(skills_text, short=1.6, medium=2.0, long=2.6),
    )

  if work_items:
    _add_section_heading(doc, "Work Experience")
  for item in work_items:
    heading = " | ".join([part for part in [item["role"], item["company"]] if part]).strip(" |")
    bullets = item["bullets"]
    entry_text = " ".join([heading, *bullets]).strip()
    entry_gap = _content_spacing(entry_text, short=1.0, medium=1.6, long=2.2)
    if heading:
      _add_subheading(doc, heading, space_after=0.4 if item["start"] or item["end"] else entry_gap)
    dates = " - ".join([part for part in [item["start"], item["end"]] if part]).strip(" -")
    if dates:
      _add_paragraph(
        doc,
        dates,
        size=META_SIZE,
        italic=True,
        color=(75, 85, 99),
        space_after=0.8 if bullets else entry_gap,
      )
    for index, bullet in enumerate(bullets):
      is_last = index == len(bullets) - 1
      bullet_row = doc.add_paragraph(style="List Bullet")
      bullet_row.paragraph_format.space_before = Pt(0)
      bullet_space = _content_spacing(bullet, short=0.8, medium=1.1, long=1.4)
      bullet_row.paragraph_format.space_after = Pt(bullet_space + (entry_gap if is_last else 0))
      bullet_run = bullet_row.add_run(bullet)
      bullet_run.font.size = Pt(BODY_SIZE)
      bullet_run.font.color.rgb = RGBColor(17, 24, 39)

  if education_items:
    _add_section_heading(doc, "Education")
  for item in education_items:
    line = " - ".join([part for part in [item["degree"], item["institution"]] if part]).strip(" -")
    if line:
      _add_subheading(doc, line, space_after=_content_spacing(line, short=1.2, medium=1.8, long=2.3))

  if project_items:
    _add_section_heading(doc, "Projects")
  for item in project_items:
    bullets = item["bullets"]
    project_text = " ".join([item["name"], item["description"], *bullets]).strip()
    project_gap = _content_spacing(project_text, short=1.0, medium=1.6, long=2.2)
    if item["name"]:
      _add_subheading(
        doc,
        item["name"],
        space_after=0.4 if item["description"] or item["link"] or bullets else project_gap,
      )
    if item["description"]:
      _add_paragraph(
        doc,
        item["description"],
        size=BODY_SIZE,
        color=(17, 24, 39),
        space_after=0.4 if item["link"] or bullets else project_gap,
      )
    if item["link"]:
      _add_paragraph(
        doc,
        item["link"],
        size=META_SIZE,
        color=(3, 105, 161),
        space_after=0.5 if bullets else project_gap,
      )
    for index, bullet in enumerate(bullets):
      is_last = index == len(bullets) - 1
      bullet_row = doc.add_paragraph(style="List Bullet")
      bullet_row.paragraph_format.space_before = Pt(0)
      bullet_space = _content_spacing(bullet, short=0.8, medium=1.1, long=1.4)
      bullet_row.paragraph_format.space_after = Pt(bullet_space + (project_gap if is_last else 0))
      bullet_run = bullet_row.add_run(bullet)
      bullet_run.font.size = Pt(BODY_SIZE)
      bullet_run.font.color.rgb = RGBColor(17, 24, 39)

  if certification_items:
    _add_section_heading(doc, "Certifications")
  for item in certification_items:
    line = " - ".join([part for part in [item["name"], item["issuer"]] if part]).strip(" -")
    if item["date"]:
      line = f"{line} ({item['date']})" if line else item["date"]
    cert_text = " ".join([line, item["link"]]).strip()
    cert_gap = _content_spacing(cert_text, short=1.0, medium=1.5, long=2.0)
    if line:
      _add_subheading(doc, line.strip(), space_after=0.4 if item["link"] else cert_gap)
    if item["link"]:
      _add_paragraph(doc, item["link"], size=META_SIZE, color=(3, 105, 161), space_after=cert_gap)

  output = BytesIO()
  doc.save(output)
  output.seek(0)
  return output.read()
